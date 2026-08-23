import os
file_name = os.path.basename(__file__)
print(f"The filename of this script is: {file_name}")

import argparse
import hashlib
import logging
import os
import re
import sys
from io import BytesIO
from pathlib import Path
from typing import Optional
from urllib.parse import unquote, urlparse

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from markdown_it import MarkdownIt
from markdown_it.token import Token
from PIL import Image, ImageOps


DEBUG = False

DEFAULT_INPUT_DIRECTORY = "markdown"
DEFAULT_OUTPUT_DIRECTORY = "."
DEFAULT_ASSET_DIRECTORY = "assets"
USER_PROFILE = Path(os.environ.get("USERPROFILE") or Path.home())
ARCHIVE_ROOT = USER_PROFILE / "Documents" / "ChatGPT Archive"

ROLE_HEADINGS = {
    "User": "Utilisateur",
    "Assistant": "Assistant",
    "System": "Système",
    "Tool": "Outil",
}

IMAGE_SUFFIXES = {
    ".png",
    ".jpg",
    ".jpeg",
    ".gif",
    ".bmp",
    ".tif",
    ".tiff",
    ".webp",
}

XML_INVALID_CONTROL_CHARACTERS = re.compile(
    r"[^\x09\x0A\x0D\x20-\uD7FF\uE000-\uFFFD\U00010000-\U0010FFFF]"
)

_SANDBOX_ASSET_NAME_CACHE: dict[str, dict[str, list[Path]]] = {}


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def choose_identical_sandbox_candidate(
    candidates: list[Path],
    asset_root: Path,
) -> Optional[Path]:
    """Return one candidate only when all same-name files are byte-identical."""
    try:
        digests = {sha256_file(path) for path in candidates}
    except OSError as exc:
        logging.warning(
            "Unable to compare ambiguous sandbox-link candidates: %s",
            exc,
        )
        return None

    if len(digests) != 1:
        return None

    def preference(path: Path) -> tuple[int, int, str]:
        try:
            relative = path.resolve().relative_to(asset_root.resolve())
            parts = relative.parts
            bucket = parts[0].lower() if parts else ""
            text = relative.as_posix()
        except ValueError:
            bucket = ""
            text = str(path)

        bucket_rank = {
            "attachment": 0,
            "image": 1,
            "dictation": 2,
        }.get(bucket, 3)
        return (bucket_rank, len(text), text.casefold())

    return min(candidates, key=preference)


def xml_safe_text(text: str) -> str:
    """Remove characters forbidden by XML 1.0 / DOCX."""
    return XML_INVALID_CONTROL_CHARACTERS.sub(
        "",
        text,
    )


def sanitize_markdown_tokens(tokens: list[object]) -> int:
    """Sanitize text produced by Markdown parsing, including decoded entities.

    The source Markdown can be XML-safe while still containing character
    references such as ``&#x0;``. markdown-it decodes those references into
    actual control characters in ``Token.content``. Sanitizing only the raw
    Markdown therefore is not sufficient before passing text to python-docx.
    """
    removed = 0

    for token in tokens:
        content = getattr(token, "content", None)
        if isinstance(content, str) and content:
            clean_content = xml_safe_text(content)
            removed += len(content) - len(clean_content)
            token.content = clean_content

        attrs = getattr(token, "attrs", None)
        if isinstance(attrs, dict):
            for key, value in list(attrs.items()):
                if isinstance(value, str):
                    clean_value = xml_safe_text(value)
                    removed += len(value) - len(clean_value)
                    attrs[key] = clean_value

        children = getattr(token, "children", None)
        if isinstance(children, list) and children:
            removed += sanitize_markdown_tokens(children)

    return removed


def configure_logging() -> None:
    level = logging.DEBUG if DEBUG else logging.INFO

    logging.basicConfig(
        level=level,
        format="%(asctime)s | %(levelname)-8s | %(message)s",
        datefmt="%H:%M:%S",
        force=True,
    )

    root_logger = logging.getLogger()
    root_logger.setLevel(level)

    # markdown-it-py emits extremely verbose parser traces at DEBUG level.
    # Force its loggers to WARNING even when this module is imported after
    # another script that previously enabled DEBUG logging.
    for logger_name in (
        "markdown_it",
        "markdown_it.main",
        "markdown_it.parser_block",
        "markdown_it.parser_inline",
        "markdown_it.rules_block",
        "markdown_it.rules_inline",
    ):
        logging.getLogger(logger_name).setLevel(logging.WARNING)


def resolve_path(
    path: Path,
    script_directory: Path,
) -> Path:
    if path.is_absolute():
        return path.resolve()

    return (
        script_directory
        / path
    ).resolve()


def ensure_styles(document: Document) -> None:
    styles = document.styles

    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    for name, size, bold in (
        ("Title", 22, True),
        ("Heading 1", 17, True),
        ("Heading 2", 14, True),
        ("Heading 3", 12, True),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = bold

    if "Chat Role" not in styles:
        style = styles.add_style(
            "Chat Role",
            WD_STYLE_TYPE.PARAGRAPH,
        )
        style.font.name = "Aptos Display"
        style.font.size = Pt(11)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)

    if "Chat Code" not in styles:
        style = styles.add_style(
            "Chat Code",
            WD_STYLE_TYPE.PARAGRAPH,
        )
        style.font.name = "Cascadia Mono"
        style.font.size = Pt(8.5)
        style.paragraph_format.left_indent = Cm(0.5)
        style.paragraph_format.right_indent = Cm(0.5)
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(6)

    if "Chat Quote" not in styles:
        style = styles.add_style(
            "Chat Quote",
            WD_STYLE_TYPE.PARAGRAPH,
        )
        style.font.name = "Aptos"
        style.font.size = Pt(10)
        style.font.italic = True
        style.paragraph_format.left_indent = Cm(0.8)
        style.paragraph_format.space_before = Pt(3)
        style.paragraph_format.space_after = Pt(3)

    if "Chat Attachment" not in styles:
        style = styles.add_style(
            "Chat Attachment",
            WD_STYLE_TYPE.PARAGRAPH,
        )
        style.font.name = "Aptos"
        style.font.size = Pt(9.5)
        style.font.italic = True
        style.paragraph_format.left_indent = Cm(0.5)
        style.paragraph_format.space_before = Pt(2)
        style.paragraph_format.space_after = Pt(4)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    ensure_styles(document)


def set_cell_shading(
    cell,
    fill: str,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))

    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)

    shading.set(qn("w:fill"), fill)


def set_paragraph_shading(
    paragraph,
    fill: str,
) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = p_pr.find(qn("w:shd"))

    if shading is None:
        shading = OxmlElement("w:shd")
        p_pr.append(shading)

    shading.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    tr_pr.append(table_header)


def add_hyperlink(
    paragraph,
    text: str,
    url: str,
) -> None:
    text = xml_safe_text(text)
    url = xml_safe_text(url)
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/"
        "officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    run.append(run_properties)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def build_sandbox_asset_name_index(
    asset_root: Path,
) -> dict[str, list[Path]]:
    cache_key = str(asset_root.resolve())
    cached = _SANDBOX_ASSET_NAME_CACHE.get(
        cache_key
    )

    if cached is not None:
        return cached

    index: dict[str, list[Path]] = {}

    if asset_root.is_dir():
        for path in sorted(asset_root.rglob("*")):
            if not path.is_file():
                continue

            aliases = {path.name}

            if "__" in path.name:
                aliases.add(
                    path.name.split("__", 1)[1]
                )

            for alias in aliases:
                index.setdefault(
                    alias.casefold(),
                    [],
                ).append(path)

    _SANDBOX_ASSET_NAME_CACHE[
        cache_key
    ] = index
    return index


def resolve_sandbox_asset_target(
    target: str,
    output_path: Path,
) -> Optional[Path]:
    parsed = urlparse(target)

    if parsed.scheme != "sandbox":
        return None

    basename = Path(
        unquote(parsed.path)
    ).name

    if not basename:
        logging.warning(
            "Sandbox link has no filename: %s",
            target,
        )
        return None

    asset_root = (
        output_path.parent
        / DEFAULT_ASSET_DIRECTORY
    )
    index = build_sandbox_asset_name_index(
        asset_root
    )
    candidates = index.get(
        basename.casefold(),
        [],
    )

    if len(candidates) == 1:
        logging.info(
            "Resolved sandbox link to archived asset: %s -> %s",
            target,
            candidates[0],
        )
        return candidates[0].resolve()

    if not candidates:
        logging.warning(
            "Sandbox link not found in archived assets; leaving text non-clickable: %s",
            target,
        )
        return None

    identical = choose_identical_sandbox_candidate(
        candidates,
        asset_root,
    )
    if identical is not None:
        logging.info(
            "Resolved sandbox link across %d byte-identical archived assets: %s -> %s",
            len(candidates),
            target,
            identical,
        )
        return identical.resolve()

    logging.warning(
        "Sandbox link is ambiguous across %d archived assets; leaving text non-clickable: %s",
        len(candidates),
        target,
    )
    return None


def normalize_local_target(
    target: str,
    markdown_path: Path,
) -> Optional[Path]:
    parsed = urlparse(target)

    if parsed.scheme in {
        "http",
        "https",
        "mailto",
        "sandbox",
    }:
        return None

    if parsed.scheme == "file":
        raw_path = unquote(parsed.path)

        if (
            os.name == "nt"
            and raw_path.startswith("/")
            and len(raw_path) >= 3
            and raw_path[2] == ":"
        ):
            raw_path = raw_path[1:]

        return Path(raw_path).resolve()

    clean_target = unquote(
        target.split("#", 1)[0]
    )

    if not clean_target:
        return None

    target_path = Path(
        clean_target.replace("/", os.sep)
    )

    if target_path.is_absolute():
        return target_path.resolve()

    return (
        markdown_path.parent
        / target_path
    ).resolve()


def windows_relative_hyperlink_target(
    relative: str,
) -> str:
    target = relative.replace("/", "\\")

    if os.sep != "\\":
        target = target.replace(os.sep, "\\")

    if not target.startswith(
        (".\\", "..\\")
    ):
        target = f".\\{target}"

    return target


def docx_hyperlink_target(
    target: str,
    markdown_path: Path,
    output_path: Path,
) -> Optional[str]:
    parsed = urlparse(target)

    if parsed.scheme in {
        "http",
        "https",
        "mailto",
    }:
        return target

    if parsed.scheme == "sandbox":
        local_path = resolve_sandbox_asset_target(
            target,
            output_path,
        )
    else:
        local_path = normalize_local_target(
            target,
            markdown_path,
        )

    if local_path is None:
        return None

    try:
        relative = os.path.relpath(
            local_path,
            output_path.parent.resolve(),
        )
    except ValueError:
        # Different Windows drives cannot be represented as a relative path.
        # Keep the absolute file URI as a last-resort fallback.
        return local_path.as_uri()

    # Atlantis Word Processor accepts Windows-style relative file targets but
    # does not resolve forward-slash or percent-encoded-space variants.  Keep
    # spaces literal and make the relative base explicit with .\ .
    return windows_relative_hyperlink_target(
        relative
    )


def is_asset_reference_paragraph(
    inline_token: Optional[Token],
) -> bool:
    if inline_token is None:
        return False

    text = paragraph_text_from_inline(
        inline_token
    )

    return text.startswith(
        (
            "📎 Archived attachment:",
            "🎵 Archived audio:",
            "🖼 Archived image:",
            "Asset:",
            "Asset ID:",
            "Archive path:",
        )
    )


def available_width_inches(document: Document) -> float:
    section = document.sections[-1]

    return (
        section.page_width.inches
        - section.left_margin.inches
        - section.right_margin.inches
    )


def image_width_inches(
    path: Path,
    document: Document,
) -> float:
    maximum = min(
        available_width_inches(document),
        6.5,
    )

    try:
        with Image.open(path) as image:
            width_pixels, height_pixels = image.size

            dpi = image.info.get("dpi", (96, 96))
            dpi_x = (
                dpi[0]
                if isinstance(dpi, tuple)
                and dpi
                else 96
            )

            if not isinstance(
                dpi_x,
                (int, float),
            ) or dpi_x <= 0:
                dpi_x = 96

            natural_width = width_pixels / dpi_x

            if height_pixels <= 0:
                return maximum

            return min(
                max(natural_width, 1.0),
                maximum,
            )

    except Exception:
        return maximum


def normalized_png_stream(
    image_path: Path,
) -> BytesIO:
    """Return a Pillow-normalized PNG stream suitable for python-docx."""
    stream = BytesIO()

    with Image.open(image_path) as source_image:
        # Apply EXIF orientation before stripping container-specific metadata.
        image = ImageOps.exif_transpose(source_image)

        # DOCX embedding only needs one raster frame. ChatGPT archive assets
        # are expected to be static images, but taking frame 0 also makes
        # animated WebP/GIF inputs deterministic.
        try:
            image.seek(0)
        except (AttributeError, EOFError):
            pass

        has_alpha = (
            image.mode in {
                "RGBA",
                "LA",
            }
            or "transparency" in image.info
        )

        image = image.convert(
            "RGBA" if has_alpha else "RGB"
        )

        save_options: dict[str, object] = {}
        dpi = source_image.info.get("dpi")

        if (
            isinstance(dpi, tuple)
            and len(dpi) >= 2
            and all(
                isinstance(value, (int, float))
                and value > 0
                for value in dpi[:2]
            )
        ):
            save_options["dpi"] = (
                float(dpi[0]),
                float(dpi[1]),
            )

        image.save(
            stream,
            format="PNG",
            **save_options,
        )

    stream.seek(0)
    return stream


def add_picture_with_fallback(
    run,
    image_path: Path,
    width_inches: float,
):
    """Embed an image directly, then normalize through Pillow if needed."""
    try:
        return run.add_picture(
            str(image_path),
            width=Inches(width_inches),
        )
    except Exception as direct_exc:
        logging.info(
            "Normalizing image for DOCX embedding: %s "
            "(%s)",
            image_path,
            type(direct_exc).__name__,
        )

        normalized_stream = normalized_png_stream(
            image_path
        )

        return run.add_picture(
            normalized_stream,
            width=Inches(width_inches),
        )


def add_image(
    document: Document,
    image_path: Path,
    alt_text: str,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = (
        WD_PARAGRAPH_ALIGNMENT.CENTER
    )

    try:
        width = image_width_inches(
            image_path,
            document,
        )

        run = paragraph.add_run()
        inline_shape = add_picture_with_fallback(
            run,
            image_path,
            width,
        )

        try:
            doc_properties = (
                inline_shape._inline.docPr
            )
            doc_properties.set(
                "descr",
                xml_safe_text(alt_text or image_path.name),
            )
        except Exception:
            pass

        if alt_text:
            caption = document.add_paragraph(
                style="Caption"
            )
            caption.alignment = (
                WD_PARAGRAPH_ALIGNMENT.CENTER
            )
            caption.add_run(xml_safe_text(alt_text))

    except Exception as exc:
        logging.warning(
            "Unable to embed image %s: %s: %s",
            image_path,
            type(exc).__name__,
            exc,
        )
        paragraph.style = "Chat Attachment"
        paragraph.add_run(
            f"[Image unavailable: "
            f"{image_path.name}]"
        )


def add_inline_tokens(
    paragraph,
    tokens: list[Token],
    markdown_path: Path,
) -> None:
    formatting_stack: list[str] = []

    for token in tokens:
        token_type = token.type

        if token_type in {
            "strong_open",
            "em_open",
            "s_open",
        }:
            formatting_stack.append(token_type)
            continue

        if token_type in {
            "strong_close",
            "em_close",
            "s_close",
        }:
            matching_open = (
                token_type.replace(
                    "_close",
                    "_open",
                )
            )

            for index in range(
                len(formatting_stack) - 1,
                -1,
                -1,
            ):
                if (
                    formatting_stack[index]
                    == matching_open
                ):
                    formatting_stack.pop(index)
                    break

            continue

        if token_type == "softbreak":
            paragraph.add_run(" ")
            continue

        if token_type == "hardbreak":
            paragraph.add_run().add_break()
            continue

        if token_type == "code_inline":
            run = paragraph.add_run(xml_safe_text(token.content))
            run.font.name = "Cascadia Mono"
            run.font.size = Pt(9)
            set_paragraph_shading(
                paragraph,
                "F2F2F2",
            )
            continue

        if token_type == "link_open":
            target = token.attrGet("href") or ""
            label_parts: list[str] = []

            continue

        if token_type == "image":
            target = token.attrGet("src") or ""
            alt_text = token.content or "image"
            local_path = normalize_local_target(
                target,
                markdown_path,
            )

            if (
                local_path is not None
                and local_path.is_file()
                and local_path.suffix.lower()
                in IMAGE_SUFFIXES
            ):
                paragraph.add_run(
                    f"[Image: {alt_text}]"
                )
            else:
                add_hyperlink(
                    paragraph,
                    alt_text,
                    target,
                )
            continue

        if token_type == "html_inline":
            paragraph.add_run(xml_safe_text(token.content))
            continue

        if token_type != "text":
            continue

        run = paragraph.add_run(xml_safe_text(token.content))

        if "strong_open" in formatting_stack:
            run.bold = True

        if "em_open" in formatting_stack:
            run.italic = True

        if "s_open" in formatting_stack:
            run.font.strike = True


def extract_link_text(
    inline_children: list[Token],
    start_index: int,
) -> tuple[str, int]:
    parts: list[str] = []
    depth = 1
    index = start_index + 1

    while index < len(inline_children):
        token = inline_children[index]

        if token.type == "link_open":
            depth += 1

        elif token.type == "link_close":
            depth -= 1

            if depth == 0:
                return "".join(parts), index

        elif token.type in {
            "text",
            "code_inline",
        }:
            parts.append(token.content)

        elif token.type in {
            "softbreak",
            "hardbreak",
        }:
            parts.append(" ")

        index += 1

    return "".join(parts), start_index


def add_inline_content(
    paragraph,
    inline_token: Token,
    markdown_path: Path,
    output_path: Path,
) -> list[tuple[str, str]]:
    deferred_images: list[tuple[str, str]] = []
    children = inline_token.children or []
    formatting_stack: list[str] = []
    index = 0

    while index < len(children):
        token = children[index]
        token_type = token.type

        if token_type in {
            "strong_open",
            "em_open",
            "s_open",
        }:
            formatting_stack.append(token_type)
            index += 1
            continue

        if token_type in {
            "strong_close",
            "em_close",
            "s_close",
        }:
            matching_open = token_type.replace(
                "_close",
                "_open",
            )

            for stack_index in range(
                len(formatting_stack) - 1,
                -1,
                -1,
            ):
                if (
                    formatting_stack[stack_index]
                    == matching_open
                ):
                    formatting_stack.pop(
                        stack_index
                    )
                    break

            index += 1
            continue

        if token_type == "link_open":
            target = token.attrGet("href") or ""
            label, closing_index = extract_link_text(
                children,
                index,
            )
            hyperlink_target = docx_hyperlink_target(
                target,
                markdown_path,
                output_path,
            )

            if hyperlink_target is None:
                paragraph.add_run(
                    xml_safe_text(label or target)
                )
            else:
                add_hyperlink(
                    paragraph,
                    label or target,
                    hyperlink_target,
                )
            index = closing_index + 1
            continue

        if token_type == "image":
            target = token.attrGet("src") or ""
            alt_text = token.content or "image"
            deferred_images.append(
                (
                    target,
                    alt_text,
                )
            )
            index += 1
            continue

        if token_type == "softbreak":
            paragraph.add_run(" ")
            index += 1
            continue

        if token_type == "hardbreak":
            paragraph.add_run().add_break()
            index += 1
            continue

        if token_type == "code_inline":
            run = paragraph.add_run(xml_safe_text(token.content))
            run.font.name = "Cascadia Mono"
            run.font.size = Pt(9)
            index += 1
            continue

        if token_type == "html_inline":
            paragraph.add_run(xml_safe_text(token.content))
            index += 1
            continue

        if token_type == "text":
            run = paragraph.add_run(xml_safe_text(token.content))

            if "strong_open" in formatting_stack:
                run.bold = True

            if "em_open" in formatting_stack:
                run.italic = True

            if "s_open" in formatting_stack:
                run.font.strike = True

        index += 1

    return deferred_images


def add_deferred_images(
    document: Document,
    images: list[tuple[str, str]],
    markdown_path: Path,
) -> None:
    for target, alt_text in images:
        local_path = normalize_local_target(
            target,
            markdown_path,
        )

        if (
            local_path is not None
            and local_path.is_file()
            and local_path.suffix.lower()
            in IMAGE_SUFFIXES
        ):
            add_image(
                document,
                local_path,
                alt_text,
            )
            continue

        paragraph = document.add_paragraph(
            style="Chat Attachment"
        )

        if target.startswith(
            (
                "http://",
                "https://",
            )
        ):
            paragraph.add_run("Image externe : ")
            add_hyperlink(
                paragraph,
                alt_text or target,
                target,
            )
        else:
            paragraph.add_run(
                f"[Image locale introuvable : "
                f"{target}]"
            )


def paragraph_text_from_inline(
    token: Token,
) -> str:
    children = token.children or []
    parts: list[str] = []

    for child in children:
        if child.type in {
            "text",
            "code_inline",
        }:
            parts.append(child.content)

        elif child.type in {
            "softbreak",
            "hardbreak",
        }:
            parts.append(" ")

    return "".join(parts).strip()


def is_role_heading(text: str) -> bool:
    normalized = text.strip().rstrip(":")

    return normalized in {
        "User",
        "Assistant",
        "System",
        "Tool",
        "Utilisateur",
        "Système",
        "Outil",
    }


def normalize_role_heading(text: str) -> str:
    normalized = text.strip().rstrip(":")

    return ROLE_HEADINGS.get(
        normalized,
        normalized,
    )


def add_role_heading(
    document: Document,
    text: str,
) -> None:
    paragraph = document.add_paragraph(
        style="Chat Role"
    )

    label = normalize_role_heading(text)
    run = paragraph.add_run(label)
    run.bold = True

    shading = {
        "Utilisateur": "EAF2F8",
        "Assistant": "E8F5E9",
        "Système": "F2F2F2",
        "Outil": "FFF4E5",
    }.get(label, "F2F2F2")

    set_paragraph_shading(
        paragraph,
        shading,
    )


def add_code_block(
    document: Document,
    content: str,
    language: str,
) -> None:
    if language:
        label = document.add_paragraph(
            style="Chat Attachment"
        )
        label.add_run(language).bold = True

    paragraph = document.add_paragraph(
        style="Chat Code"
    )
    set_paragraph_shading(
        paragraph,
        "F3F3F3",
    )

    run = paragraph.add_run(
        content.rstrip("\n")
    )
    run.font.name = "Cascadia Mono"
    run.font.size = Pt(8.5)


def parse_table(
    tokens: list[Token],
    start_index: int,
) -> tuple[list[list[Token]], int]:
    rows: list[list[Token]] = []
    current_row: list[Token] = []
    index = start_index + 1

    while index < len(tokens):
        token = tokens[index]

        if token.type == "table_close":
            if current_row:
                rows.append(current_row)
            return rows, index

        if token.type == "tr_open":
            current_row = []

        elif token.type == "tr_close":
            rows.append(current_row)
            current_row = []

        elif token.type in {
            "th_open",
            "td_open",
        }:
            if (
                index + 1 < len(tokens)
                and tokens[index + 1].type
                == "inline"
            ):
                current_row.append(
                    tokens[index + 1]
                )

        index += 1

    return rows, start_index


def add_table(
    document: Document,
    rows: list[list[Token]],
    markdown_path: Path,
    output_path: Path,
) -> None:
    if not rows:
        return

    column_count = max(
        len(row)
        for row in rows
    )

    table = document.add_table(
        rows=len(rows),
        cols=column_count,
    )
    table.style = "Table Grid"

    for row_index, row in enumerate(rows):
        word_row = table.rows[row_index]

        if row_index == 0:
            set_repeat_table_header(word_row)

        for column_index in range(column_count):
            cell = word_row.cells[column_index]
            cell.vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.TOP
            )

            if row_index == 0:
                set_cell_shading(
                    cell,
                    "D9EAF7",
                )

            paragraph = cell.paragraphs[0]

            if column_index < len(row):
                deferred = add_inline_content(
                    paragraph,
                    row[column_index],
                    markdown_path,
                    output_path,
                )

                for target, alt_text in deferred:
                    local_path = normalize_local_target(
                        target,
                        markdown_path,
                    )

                    if (
                        local_path is not None
                        and local_path.is_file()
                        and local_path.suffix.lower()
                        in IMAGE_SUFFIXES
                    ):
                        try:
                            run = paragraph.add_run()
                            add_picture_with_fallback(
                                run,
                                local_path,
                                1.5,
                            )
                        except Exception as exc:
                            logging.warning(
                                "Unable to embed table image %s: %s: %s",
                                local_path,
                                type(exc).__name__,
                                exc,
                            )
                            paragraph.add_run(
                                f"[{alt_text}]"
                            )


def add_horizontal_rule(
    document: Document,
) -> None:
    paragraph = document.add_paragraph()
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B7B7B7")
    borders.append(bottom)
    p_pr.append(borders)


def convert_markdown_to_docx(
    markdown_path: Path,
    output_path: Path,
    document_title: Optional[str],
) -> None:
    print(
        f"Converting: {markdown_path.name}",
        flush=True,
    )
    logging.info(
        "Reading Markdown: %s",
        markdown_path,
    )

    markdown_text = markdown_path.read_text(
        encoding="utf-8",
        errors="strict",
    )

    clean_markdown_text = xml_safe_text(markdown_text)

    removed_character_count = (
        len(markdown_text)
        - len(clean_markdown_text)
    )

    if removed_character_count:
        logging.warning(
            "Removed %d XML-incompatible control "
            "character(s) from %s",
            removed_character_count,
            markdown_path.name,
        )

    markdown_text = clean_markdown_text

    markdown = (
        MarkdownIt(
            "commonmark",
            {
                "html": True,
                "linkify": True,
                "typographer": False,
            },
        )
        .enable("table")
        .enable("strikethrough")
    )

    tokens = markdown.parse(markdown_text)

    parsed_removed_character_count = sanitize_markdown_tokens(tokens)
    if parsed_removed_character_count:
        logging.warning(
            "Removed %d XML-incompatible character(s) produced by Markdown parsing from %s",
            parsed_removed_character_count,
            markdown_path.name,
        )

    document = Document()
    configure_document(document)

    if document_title:
        title = document.add_paragraph(
            style="Title"
        )
        title.alignment = (
            WD_PARAGRAPH_ALIGNMENT.CENTER
        )
        title.add_run(xml_safe_text(document_title))

    list_stack: list[dict[str, int | bool]] = []
    blockquote_depth = 0
    index = 0

    while index < len(tokens):
        token = tokens[index]
        token_type = token.type

        if token_type == "heading_open":
            level = int(
                token.tag[1:]
                if token.tag.startswith("h")
                else "1"
            )

            inline = (
                tokens[index + 1]
                if index + 1 < len(tokens)
                else None
            )

            text = (
                paragraph_text_from_inline(inline)
                if inline is not None
                and inline.type == "inline"
                else ""
            )

            if is_role_heading(text):
                add_role_heading(
                    document,
                    text,
                )
            else:
                paragraph = document.add_paragraph(
                    style=f"Heading {min(level, 3)}"
                )

                if inline is not None:
                    deferred = add_inline_content(
                        paragraph,
                        inline,
                        markdown_path,
                        output_path,
                    )
                    add_deferred_images(
                        document,
                        deferred,
                        markdown_path,
                    )

            index += 3
            continue

        if token_type == "paragraph_open":
            inline = (
                tokens[index + 1]
                if index + 1 < len(tokens)
                else None
            )

            style = None

            if blockquote_depth > 0:
                style = "Chat Quote"
            elif is_asset_reference_paragraph(inline):
                style = "Chat Attachment"

            paragraph = document.add_paragraph(
                style=style
            )

            if list_stack:
                current = list_stack[-1]
                ordered = bool(
                    current["ordered"]
                )
                level = len(list_stack)

                paragraph.style = (
                    "List Number"
                    if ordered
                    else "List Bullet"
                )
                paragraph.paragraph_format.left_indent = Cm(
                    0.63 * level
                )

            if inline is not None:
                deferred = add_inline_content(
                    paragraph,
                    inline,
                    markdown_path,
                    output_path,
                )
                add_deferred_images(
                    document,
                    deferred,
                    markdown_path,
                )

            index += 3
            continue

        if token_type in {
            "fence",
            "code_block",
        }:
            add_code_block(
                document,
                token.content,
                token.info.strip(),
            )
            index += 1
            continue

        if token_type == "blockquote_open":
            blockquote_depth += 1
            index += 1
            continue

        if token_type == "blockquote_close":
            blockquote_depth = max(
                blockquote_depth - 1,
                0,
            )
            index += 1
            continue

        if token_type == "bullet_list_open":
            list_stack.append(
                {
                    "ordered": False,
                    "start": 1,
                }
            )
            index += 1
            continue

        if token_type == "ordered_list_open":
            start = token.attrGet("start")

            list_stack.append(
                {
                    "ordered": True,
                    "start": (
                        int(start)
                        if start
                        else 1
                    ),
                }
            )
            index += 1
            continue

        if token_type in {
            "bullet_list_close",
            "ordered_list_close",
        }:
            if list_stack:
                list_stack.pop()
            index += 1
            continue

        if token_type == "hr":
            add_horizontal_rule(document)
            index += 1
            continue

        if token_type == "table_open":
            rows, closing_index = parse_table(
                tokens,
                index,
            )
            add_table(
                document,
                rows,
                markdown_path,
                output_path,
            )
            index = closing_index + 1
            continue

        if token_type == "html_block":
            paragraph = document.add_paragraph()
            paragraph.add_run(
                token.content.strip()
            )
            index += 1
            continue

        index += 1

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    core_properties = document.core_properties
    core_properties.title = (
        document_title
        or markdown_path.stem
    )
    core_properties.subject = (
        "ChatGPT conversation export"
    )
    core_properties.keywords = (
        "ChatGPT, conversation, export"
    )

    print(
        f"Saving:     {output_path}",
        flush=True,
    )
    logging.info(
        "Writing DOCX: %s",
        output_path,
    )

    document.save(output_path)

    size = output_path.stat().st_size
    print(
        f"Created:    {output_path.name} "
        f"({size:,} bytes)",
        flush=True,
    )
    logging.info(
        "DOCX written: %d bytes",
        size,
    )


def parse_arguments(
    script_directory: Path,
) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Convert a ChatGPT Markdown export to DOCX "
            "with local images embedded."
        )
    )

    parser.add_argument(
        "input",
        type=Path,
        nargs="?",
        default=None,
        help=(
            "Markdown input file. If omitted, convert "
            "all .md files from the archive markdown directory."
        ),
    )

    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        default=None,
        help=(
            "Exact DOCX output path. Valid only with "
            "one input file."
        ),
    )

    parser.add_argument(
        "--input-directory",
        type=Path,
        default=(
            ARCHIVE_ROOT
            / DEFAULT_INPUT_DIRECTORY
        ),
    )

    parser.add_argument(
        "--output-directory",
        type=Path,
        default=(
            ARCHIVE_ROOT
            / DEFAULT_OUTPUT_DIRECTORY
        ),
    )

    parser.add_argument(
        "--title",
        type=str,
        default=None,
        help=(
            "Optional document title for a single file."
        ),
    )

    parser.add_argument(
        "--quiet",
        action="store_true",
    )

    parser.add_argument(
        "--overwrite",
        action="store_true",
        help=(
            "Regenerate DOCX files even when a non-empty "
            "output file already exists."
        ),
    )

    return parser.parse_args()


def main() -> int:
    script_directory = Path(__file__).resolve().parent
    arguments = parse_arguments(script_directory)

    configure_logging()

    if arguments.quiet:
        logging.getLogger().setLevel(logging.INFO)

    try:
        output_directory = resolve_path(
            arguments.output_directory,
            script_directory,
        )

        if arguments.input is not None:
            input_path = resolve_path(
                arguments.input,
                script_directory,
            )

            if not input_path.exists():
                raise FileNotFoundError(
                    f"Markdown file not found: {input_path}"
                )

            if arguments.output is None:
                output_path = (
                    output_directory
                    / input_path.with_suffix(
                        ".docx"
                    ).name
                )
            else:
                output_path = resolve_path(
                    arguments.output,
                    script_directory,
                )

            if (
                not arguments.overwrite
                and output_path.is_file()
                and output_path.stat().st_size > 0
            ):
                print(
                    f"SKIP existing DOCX: {output_path}"
                )
                return 0

            convert_markdown_to_docx(
                markdown_path=input_path,
                output_path=output_path,
                document_title=arguments.title,
            )

            print()
            print(
                f"DOCX export created: "
                f"{output_path}"
            )
            return 0

        if arguments.output is not None:
            raise ValueError(
                "--output requires a single input file."
            )

        input_directory = resolve_path(
            arguments.input_directory,
            script_directory,
        )

        markdown_files = sorted(
            input_directory.glob("*.md")
        )

        if not markdown_files:
            raise FileNotFoundError(
                f"No Markdown files found in: "
                f"{input_directory}"
            )

        failures = 0
        converted = 0
        skipped = 0

        for number, markdown_path in enumerate(
            markdown_files,
            start=1,
        ):
            output_path = (
                output_directory
                / markdown_path.with_suffix(
                    ".docx"
                ).name
            )

            print(
                f"[{number}/{len(markdown_files)}] "
                f"{markdown_path.name}"
            )

            if (
                not arguments.overwrite
                and output_path.is_file()
                and output_path.stat().st_size > 0
            ):
                print(
                    f"SKIP existing DOCX: "
                    f"{output_path.name}"
                )
                skipped += 1
                continue

            try:
                convert_markdown_to_docx(
                    markdown_path=markdown_path,
                    output_path=output_path,
                    document_title=None,
                )
                converted += 1
            except Exception:
                failures += 1
                logging.exception(
                    "Failed to convert %s",
                    markdown_path,
                )

        print()
        print("DOCX batch summary")
        print("==================")
        print(
            f"Requested : {len(markdown_files)}"
        )
        print(f"Converted : {converted}")
        print(f"Skipped   : {skipped}")
        print(f"Failed    : {failures}")
        print(
            f"Output    : {output_directory}"
        )

        return 1 if failures else 0

    except KeyboardInterrupt:
        logging.error(
            "Operation cancelled by user."
        )
        return 130

    except Exception as exc:
        logging.exception(
            "DOCX exporter failed: %s",
            exc,
        )
        return 1


if __name__ == "__main__":
    sys.exit(main())
