"""Read-only scanner for historical ChatGPT conversations saved as DOCX.

This module intentionally stops before archive import.  It extracts stable
provenance, filename hints, Word metadata, document structure, and conservative
conversation-boundary hints so a corpus can be audited before any canonical
JSON/XZ or SQLite data is created.
"""

from __future__ import annotations

import datetime as dt
import hashlib
import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph


NORMALIZED_FILENAME_RE = re.compile(
    r"^(?P<category>.+?)\s+GPT\s+"
    r"(?P<date>\d{4}-\d{2}-\d{2})\s+"
    r"(?P<title>.+)\.docx$",
    re.IGNORECASE,
)
LEGACY_TIMESTAMP_FILENAME_RE = re.compile(
    r"^(?P<category>.+?)\s+GPT\s+"
    r"(?P<date>\d{4}-\d{2}-\d{2})\s+"
    r"(?P<time>\d{6})\s+"
    r"(?P<title>.+)\.docx$",
    re.IGNORECASE,
)
LEGACY_SPACED_DATE_FILENAME_RE = re.compile(
    r"^(?P<category>.+?)\s+GPT\s+"
    r"(?P<year>\d{4})\s+(?P<month>\d{2})\s+(?P<day>\d{2})\s+"
    r"(?P<title>.+)\.docx$",
    re.IGNORECASE,
)
CHATGPT_HYPERLINK_SENTINEL = 'HYPERLINK "https://chatgpt.com/"'


@dataclass(frozen=True, slots=True)
class LegacyFilenameMetadata:
    """Metadata inferred from a historical DOCX filename."""

    category_hint: str | None
    date_hint: str | None
    title_hint: str
    normalized: bool
    legacy_time_hint: str | None = None


@dataclass(frozen=True, slots=True)
class LegacyDocxReport:
    """Read-only inventory result for one historical DOCX."""

    path: str
    filename: str
    sha256: str
    size_bytes: int
    category_hint: str | None
    filename_date_hint: str | None
    filename_title_hint: str
    filename_normalized: bool
    legacy_time_hint: str | None
    docx_created_at: str | None
    docx_modified_at: str | None
    paragraph_count: int
    nonempty_paragraph_count: int
    table_count: int
    heading_count: int
    hyperlink_sentinel_count: int
    first_visible_text: str | None
    likely_first_user_message: str | None
    boundary_candidate_count: int
    parse_confidence: str
    notes: tuple[str, ...]

    def to_dict(self) -> dict[str, object]:
        """Return a JSON-serializable representation."""

        result = asdict(self)
        result["notes"] = list(self.notes)
        return result


def _clean(value: str) -> str:
    return " ".join(value.split()).strip()


def parse_legacy_filename(path: Path | str) -> LegacyFilenameMetadata:
    """Parse the supported manual filename conventions.

    Preferred convention::

        <CATEGORY> GPT <YYYY-MM-DD> <TITLE>.docx

    Two historical forms are still recognized so an audit can be run before a
    directory has been fully renamed.
    """

    filename = Path(path).name

    match = NORMALIZED_FILENAME_RE.match(filename)
    if match:
        return LegacyFilenameMetadata(
            category_hint=_clean(match.group("category")),
            date_hint=match.group("date"),
            title_hint=_clean(match.group("title")),
            normalized=True,
        )

    match = LEGACY_TIMESTAMP_FILENAME_RE.match(filename)
    if match:
        return LegacyFilenameMetadata(
            category_hint=_clean(match.group("category")),
            date_hint=match.group("date"),
            title_hint=_clean(match.group("title")),
            normalized=False,
            legacy_time_hint=match.group("time"),
        )

    match = LEGACY_SPACED_DATE_FILENAME_RE.match(filename)
    if match:
        date_hint = "-".join(
            [match.group("year"), match.group("month"), match.group("day")]
        )
        return LegacyFilenameMetadata(
            category_hint=_clean(match.group("category")),
            date_hint=date_hint,
            title_hint=_clean(match.group("title")),
            normalized=False,
        )

    return LegacyFilenameMetadata(
        category_hint=None,
        date_hint=None,
        title_hint=Path(filename).stem,
        normalized=False,
    )


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as input_file:
        for chunk in iter(lambda: input_file.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _iso_datetime(value: dt.datetime | None) -> str | None:
    if value is None:
        return None
    if value.tzinfo is None:
        value = value.replace(tzinfo=dt.timezone.utc)
    return value.isoformat()


def _paragraph_style_name(paragraph: Paragraph) -> str:
    try:
        return paragraph.style.name or ""
    except (AttributeError, KeyError):
        return ""


def _visible_paragraphs(document: DocxDocument) -> list[tuple[int, str, str]]:
    result: list[tuple[int, str, str]] = []
    for index, paragraph in enumerate(document.paragraphs):
        text = _clean(paragraph.text)
        result.append((index, text, _paragraph_style_name(paragraph)))
    return result


def _first_visible_text(paragraphs: Iterable[tuple[int, str, str]]) -> str | None:
    for _, text, _ in paragraphs:
        if text and text != CHATGPT_HYPERLINK_SENTINEL:
            return text
    return None


def _boundary_candidates(paragraphs: list[tuple[int, str, str]]) -> list[int]:
    """Return conservative possible user-turn boundaries.

    Historical browser-to-Word copies often preserve one empty paragraph
    between ordinary rendered blocks and occasionally a larger gap around a
    turn or rich-content transition.  A two-blank gap is therefore useful as
    an audit hint, but never authoritative enough to assign roles by itself.
    """

    candidates: list[int] = []
    blank_run = 0
    seen_visible = False

    for index, text, _ in paragraphs:
        if not text:
            blank_run += 1
            continue
        if text == CHATGPT_HYPERLINK_SENTINEL:
            blank_run = 0
            continue
        if not seen_visible:
            candidates.append(index)
            seen_visible = True
        elif blank_run >= 2:
            candidates.append(index)
        blank_run = 0

    return candidates


def scan_legacy_docx(path: Path | str) -> LegacyDocxReport:
    """Inspect one legacy DOCX without modifying it or the archive."""

    source = Path(path).expanduser().resolve()
    if not source.is_file():
        raise FileNotFoundError(source)
    if source.suffix.lower() != ".docx":
        raise ValueError(f"Expected a .docx file: {source}")

    filename_meta = parse_legacy_filename(source)
    document = Document(source)
    paragraphs = _visible_paragraphs(document)
    nonempty = [item for item in paragraphs if item[1]]
    headings = [
        item
        for item in nonempty
        if item[2].lower().startswith("heading")
        or item[2].lower().startswith("title")
    ]
    sentinel_count = sum(
        1 for _, text, _ in nonempty if text == CHATGPT_HYPERLINK_SENTINEL
    )
    first_text = _first_visible_text(paragraphs)
    candidates = _boundary_candidates(paragraphs)

    notes: list[str] = []
    if filename_meta.normalized:
        notes.append("filename matches preferred legacy naming convention")
    elif filename_meta.date_hint:
        notes.append("historical filename recognized; normalization recommended")
    else:
        notes.append("filename metadata could not be parsed")

    if sentinel_count:
        notes.append("ChatGPT hyperlink sentinel preserved by Word")
    if document.tables:
        notes.append("tables preserved in DOCX")
    if headings:
        notes.append("heading styles preserved in DOCX")

    # The scanner deliberately does not claim turn reconstruction yet.  The
    # first prompt is usually recoverable; later boundaries remain heuristic.
    if first_text:
        likely_first_user_message = first_text
    else:
        likely_first_user_message = None
        notes.append("no visible conversation text found")

    if first_text and (headings or document.tables or sentinel_count):
        confidence = "medium"
    elif first_text:
        confidence = "low"
    else:
        confidence = "none"

    properties = document.core_properties
    return LegacyDocxReport(
        path=str(source),
        filename=source.name,
        sha256=_sha256(source),
        size_bytes=source.stat().st_size,
        category_hint=filename_meta.category_hint,
        filename_date_hint=filename_meta.date_hint,
        filename_title_hint=filename_meta.title_hint,
        filename_normalized=filename_meta.normalized,
        legacy_time_hint=filename_meta.legacy_time_hint,
        docx_created_at=_iso_datetime(properties.created),
        docx_modified_at=_iso_datetime(properties.modified),
        paragraph_count=len(paragraphs),
        nonempty_paragraph_count=len(nonempty),
        table_count=len(document.tables),
        heading_count=len(headings),
        hyperlink_sentinel_count=sentinel_count,
        first_visible_text=first_text,
        likely_first_user_message=likely_first_user_message,
        boundary_candidate_count=len(candidates),
        parse_confidence=confidence,
        notes=tuple(notes),
    )


def scan_legacy_directory(
    directory: Path | str,
    *,
    recursive: bool = True,
) -> list[LegacyDocxReport]:
    """Scan all DOCX files below a directory in stable filename order."""

    root = Path(directory).expanduser().resolve()
    if not root.is_dir():
        raise NotADirectoryError(root)

    iterator = root.rglob("*.docx") if recursive else root.glob("*.docx")
    paths = sorted(iterator, key=lambda item: str(item).casefold())
    return [scan_legacy_docx(path) for path in paths]
