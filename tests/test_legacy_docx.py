import os
file_name = os.path.basename(__file__)
print(f"The filename of this script is: {file_name}")

import tempfile
import unittest
from pathlib import Path

from docx import Document

from gpt_exporter.legacy.docx import (
    parse_legacy_filename,
    scan_legacy_directory,
    scan_legacy_docx,
)


class LegacyDocxTests(unittest.TestCase):
    def _build_docx(self, path: Path) -> Path:
        document = Document()
        document.add_paragraph('HYPERLINK "https://chatgpt.com/"')
        document.add_paragraph("")
        document.add_paragraph("")
        document.add_paragraph("première question utilisateur")
        document.add_paragraph("")
        document.add_heading("Réponse structurée", level=2)
        document.add_paragraph("Premier paragraphe de réponse.")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Item"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "A"
        table.cell(1, 1).text = "B"
        document.add_paragraph("")
        document.add_paragraph("")
        document.add_paragraph("question suivante")
        document.add_paragraph("")
        document.add_paragraph("Deuxième réponse.")
        document.core_properties.created = None
        document.save(path)
        return path

    def test_preferred_filename_convention(self) -> None:
        metadata = parse_legacy_filename(
            "PKI GPT 2026-04-03 Installation PiOS sur RPi.docx"
        )

        self.assertEqual(metadata.category_hint, "PKI")
        self.assertEqual(metadata.date_hint, "2026-04-03")
        self.assertEqual(metadata.title_hint, "Installation PiOS sur RPi")
        self.assertTrue(metadata.normalized)
        self.assertIsNone(metadata.legacy_time_hint)

    def test_historical_timestamp_filename_is_still_recognized(self) -> None:
        metadata = parse_legacy_filename(
            "HAM GPT 2026-04-03 154435 BOM transverter 3.2-20 GHz.docx"
        )

        self.assertEqual(metadata.category_hint, "HAM")
        self.assertEqual(metadata.date_hint, "2026-04-03")
        self.assertEqual(metadata.legacy_time_hint, "154435")
        self.assertFalse(metadata.normalized)

    def test_historical_spaced_date_filename_is_still_recognized(self) -> None:
        metadata = parse_legacy_filename(
            "IT GPT 2026 03 27 Problème envoi email Gandi.docx"
        )

        self.assertEqual(metadata.category_hint, "IT")
        self.assertEqual(metadata.date_hint, "2026-03-27")
        self.assertEqual(metadata.title_hint, "Problème envoi email Gandi")
        self.assertFalse(metadata.normalized)

    def test_scanner_is_read_only_and_reports_structure(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            path = Path(temporary) / "PKI GPT 2026-04-03 Test conversation.docx"
            self._build_docx(path)
            before = path.read_bytes()

            report = scan_legacy_docx(path)

            self.assertEqual(path.read_bytes(), before)
            self.assertEqual(report.category_hint, "PKI")
            self.assertEqual(report.filename_date_hint, "2026-04-03")
            self.assertEqual(report.filename_title_hint, "Test conversation")
            self.assertEqual(report.table_count, 1)
            self.assertGreaterEqual(report.heading_count, 1)
            self.assertEqual(report.hyperlink_sentinel_count, 1)
            self.assertEqual(
                report.likely_first_user_message,
                "première question utilisateur",
            )
            self.assertGreaterEqual(report.boundary_candidate_count, 2)
            self.assertEqual(len(report.sha256), 64)
            self.assertIn(report.parse_confidence, {"medium", "low"})

    def test_directory_scan_is_stable_and_recursive(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            nested = root / "nested"
            nested.mkdir()
            self._build_docx(root / "B GPT 2026-04-03 Second.docx")
            self._build_docx(nested / "A GPT 2026-04-02 First.docx")

            reports = scan_legacy_directory(root)

            self.assertEqual(len(reports), 2)
            self.assertEqual(
                [Path(report.path).name for report in reports],
                [
                    "B GPT 2026-04-03 Second.docx",
                    "A GPT 2026-04-02 First.docx",
                ],
            )

    def test_non_docx_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            path = Path(temporary) / "not-a-docx.txt"
            path.write_text("hello", encoding="utf-8")

            with self.assertRaises(ValueError):
                scan_legacy_docx(path)


if __name__ == "__main__":
    unittest.main()
