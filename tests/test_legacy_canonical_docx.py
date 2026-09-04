import os
file_name = os.path.basename(__file__)
print(f"The filename of this script is: {file_name}")

import hashlib
import tempfile
import unittest
from pathlib import Path

from docx import Document

from gpt_exporter.legacy.canonical_docx import (
    CANONICAL_LEGACY_DOCX_VERSION,
    build_legacy_markdown,
    canonical_output_name,
    export_legacy_canonical_docx,
)


class LegacyCanonicalDocxTests(unittest.TestCase):
    def conversation(self) -> dict[str, object]:
        return {
            "source_filename": "HAM GPT 2026-03-27 Analyse avec SSA3021X part1.docx",
            "source_sha256": "a" * 64,
            "title_hint": "Analyse avec SSA3021X part1",
            "category_hint": "HAM",
            "date_hint": "2026-03-27",
            "parser_version": "legacy-docx-parser-v2",
            "role_inference_version": "legacy-role-inference-v3",
            "turn_builder_version": "legacy-turn-builder-v1",
            "starts_mid_conversation": False,
            "turns": [
                {
                    "role": "user",
                    "confidence": "high",
                    "first_order": 3,
                    "last_order": 3,
                    "content": "salut josh. on continue nos grands rangements !",
                },
                {
                    "role": "unknown",
                    "confidence": "none",
                    "first_order": 5,
                    "last_order": 8,
                    "content": "Salut 🙂 Parfait, mission grand rangement activée.",
                },
                {
                    "role": "assistant",
                    "confidence": "medium",
                    "first_order": 10,
                    "last_order": 20,
                    "content": "Parfait — on continue.",
                },
            ],
        }

    def test_markdown_preserves_provenance_unknown_turns_and_text_only_scope(self) -> None:
        markdown = build_legacy_markdown(self.conversation())
        self.assertIn("Legacy DOCX normalized derivative (text-only)", markdown)
        self.assertIn("images, embedded files, and other attachments", markdown)
        self.assertIn("Consult the source DOCX", markdown)
        self.assertIn(CANONICAL_LEGACY_DOCX_VERSION, markdown)
        self.assertNotIn("# Analyse avec SSA3021X part1", markdown)
        self.assertIn("## User", markdown)
        self.assertIn("## Unknown", markdown)
        self.assertIn("## Assistant", markdown)
        self.assertIn("1 turn(s) remain `UNKNOWN`", markdown)
        self.assertIn("grands rangements", markdown)

    def test_output_name_cannot_overwrite_historical_filename(self) -> None:
        name = canonical_output_name(self.conversation())
        self.assertTrue(name.endswith(" [normalized].docx"))
        self.assertNotEqual(name, self.conversation()["source_filename"])

    def test_export_creates_readable_docx_in_separate_directory(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            output_dir = Path(temporary) / "normalized"
            result = export_legacy_canonical_docx(
                self.conversation(),
                output_dir,
                overwrite=True,
            )
            self.assertTrue(result.output_path.is_file())
            self.assertEqual(result.turn_count, 3)
            self.assertEqual(result.unknown_turn_count, 1)
            self.assertFalse(result.source_text_restored)
            self.assertIn("[normalized]", result.output_path.name)

            document = Document(result.output_path)
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            text = "\n".join(paragraphs)
            self.assertIn("Analyse avec SSA3021X part1", text)
            self.assertEqual(
                sum(paragraph.strip() == "Analyse avec SSA3021X part1" for paragraph in paragraphs),
                1,
            )
            self.assertIn("grands rangements", text)
            self.assertIn("Unknown", text)
            self.assertIn("text-only", text)
            self.assertIn("attachments", text)

    def test_export_restores_manual_line_breaks_from_immutable_source(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            nested = root / "nested"
            nested.mkdir()
            source = nested / "HAM GPT 2026-03-27 Lines.docx"
            source_document = Document()
            paragraph = source_document.add_paragraph()
            paragraph.add_run("first code line")
            paragraph.add_run().add_break()
            paragraph.add_run("second code line")
            source_document.save(source)
            source_sha = hashlib.sha256(source.read_bytes()).hexdigest()

            conversation = {
                "source_filename": source.name,
                "source_sha256": source_sha,
                "title_hint": "Lines",
                "parser_version": "legacy-docx-parser-v2",
                "role_inference_version": "legacy-role-inference-v3",
                "turn_builder_version": "legacy-turn-builder-v1",
                "turns": [
                    {
                        "role": "assistant",
                        "confidence": "high",
                        "first_order": 0,
                        "last_order": 0,
                        "source_orders": [0],
                        "content": "first code line second code line",
                    }
                ],
            }

            result = export_legacy_canonical_docx(
                conversation,
                root / "normalized",
                overwrite=True,
                docx_root=root,
            )
            self.assertTrue(result.source_text_restored)

            document = Document(result.output_path)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("first code line", text)
            self.assertIn("second code line", text)
            self.assertIn("source Word blocks restored", text)

    def test_explicit_docx_root_refuses_missing_source(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            root = Path(temporary)
            conversation = self.conversation()
            with self.assertRaises(FileNotFoundError):
                export_legacy_canonical_docx(
                    conversation,
                    root / "normalized",
                    overwrite=True,
                    docx_root=root,
                )


if __name__ == "__main__":
    unittest.main()
