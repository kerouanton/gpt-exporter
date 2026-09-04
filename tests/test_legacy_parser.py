import os
file_name = os.path.basename(__file__)
print(f"The filename of this script is: {file_name}")

import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.shared import Inches

from gpt_exporter.legacy import LEGACY_SCHEMA, PARSER_VERSION, parse_legacy_conversation


class LegacyParserTests(unittest.TestCase):
    def _write_docx(self, path: Path, *, first_text: str) -> Path:
        document = Document()
        document.add_paragraph('HYPERLINK "https://chatgpt.com/"')
        document.add_paragraph("")
        first = document.add_paragraph()
        first.paragraph_format.left_indent = Inches(0.25)
        first.add_run(first_text).bold = True
        document.add_paragraph("")
        document.add_heading("Structured answer", level=2)
        body = document.add_paragraph("Body text")
        body.add_run(" italic").italic = True
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        document.save(path)
        return path

    def test_parser_preserves_source_and_builds_ordered_blocks(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            path = Path(temporary) / "PKI GPT 2026-04-03 Test conversation.docx"
            self._write_docx(path, first_text="question utilisateur")
            before = path.read_bytes()

            conversation = parse_legacy_conversation(path)

            self.assertEqual(path.read_bytes(), before)
            self.assertEqual(conversation.schema, LEGACY_SCHEMA)
            self.assertEqual(conversation.parser_version, PARSER_VERSION)
            self.assertEqual(conversation.source_type, "legacy_docx")
            self.assertEqual(conversation.category_hint, "PKI")
            self.assertEqual(conversation.date_hint, "2026-04-03")
            self.assertEqual(conversation.title_hint, "Test conversation")
            self.assertEqual(len(conversation.source_sha256), 64)
            self.assertGreaterEqual(len(conversation.blocks), 4)
            self.assertEqual(conversation.blocks[0].kind, "hyperlink_sentinel")
            self.assertEqual(conversation.blocks[1].kind, "paragraph")
            self.assertEqual(conversation.blocks[2].kind, "heading")
            self.assertEqual(conversation.blocks[-1].kind, "table")
            self.assertTrue(all(block.role == "unknown" for block in conversation.blocks))
            self.assertIsNone(conversation.starts_mid_conversation)

            first = conversation.blocks[1]
            self.assertEqual(first.blank_blocks_before, 1)
            self.assertIsNotNone(first.left_indent_emu)
            self.assertEqual(first.run_count, 1)
            self.assertEqual(first.bold_run_count, 1)

    def test_assistant_like_opening_marks_possible_mid_conversation_start(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            path = Path(temporary) / "HAM GPT 2026-04-23 Test continuation.docx"
            self._write_docx(path, first_text="Parfait, on voit clairement le résultat.")

            conversation = parse_legacy_conversation(path)

            self.assertTrue(conversation.starts_mid_conversation)
            self.assertEqual(conversation.starts_mid_conversation_confidence, "medium")
            self.assertTrue(
                any("assistant continuation" in note for note in conversation.notes)
            )
            self.assertTrue(all(block.role == "unknown" for block in conversation.blocks))

    def test_to_dict_is_json_ready(self) -> None:
        with tempfile.TemporaryDirectory() as temporary:
            path = Path(temporary) / "IT GPT 2026-03-27 Test.docx"
            self._write_docx(path, first_text="plain opening")

            payload = parse_legacy_conversation(path).to_dict()

            self.assertEqual(payload["schema"], LEGACY_SCHEMA)
            self.assertIsInstance(payload["blocks"], list)
            self.assertIsInstance(payload["notes"], list)
            self.assertIn("blank_blocks_before", payload["blocks"][1])
            self.assertIn("run_count", payload["blocks"][1])


if __name__ == "__main__":
    unittest.main()
